"""DOCX generation via python-docx."""

from __future__ import annotations

from pathlib import Path

from docx import Document

from app.sandbox.path_guard import normalize_user_path
from app.tools.builtin.docgen.coerce import coerce_docx_outline, docx_outline_has_body
from app.tools.builtin.docgen.schemas import DocxOutline


async def gen(outline: DocxOutline | dict, out_path: str) -> str:
    """Render *outline* to a .docx file at *out_path* and return the absolute path."""
    data = (
        outline
        if isinstance(outline, DocxOutline)
        else coerce_docx_outline(outline)
    )
    if not docx_outline_has_body(data):
        raise ValueError(
            "docx outline has title only — paragraphs[].body (or content/sections) "
            "must include non-empty text"
        )
    target = normalize_user_path(out_path)
    target.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_heading(data.title, level=0)
    for para in data.paragraphs:
        heading = (para.heading or "").strip()
        body = (para.body or "").strip()
        if heading:
            doc.add_heading(heading, level=1)
        if not body:
            continue
        for block in body.replace("\r\n", "\n").split("\n\n"):
            block = block.strip()
            if block:
                doc.add_paragraph(block)
    doc.save(str(target))
    return str(target)
