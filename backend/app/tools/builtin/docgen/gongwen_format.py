"""Apply 公文正文稿 page/font/line-spacing to an existing .docx.

Default layout (not GB/T 9704 套红):
- margins 3cm / 3cm / 2.9cm / 2.9cm
- header 1.5cm, footer 1.75cm
- exact 29pt line spacing
- 方正* GBK fonts + Times New Roman for Latin
- footer page number with long dashes, 宋体 小四
"""

from __future__ import annotations

import re
from contextvars import ContextVar
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

FONT_TITLE = "方正小标宋_GBK"
FONT_KAI = "方正楷体_GBK"
FONT_HEI = "方正黑体_GBK"
FONT_FANGSONG = "方正仿宋_GBK"
FONT_LATIN = "Times New Roman"
FONT_PAGE = "宋体"

SIZE_TITLE_PT = 22  # 二号
SIZE_BODY_PT = 16  # 三号
SIZE_PAGE_PT = 12  # 小四
LINE_PT = 29

_EAST_ASIA_REPLACE = {
    "仿宋_GB2312": FONT_FANGSONG,
    "仿宋_GBK": FONT_FANGSONG,
    "FangSong": FONT_FANGSONG,
    "FangSong_GB2312": FONT_FANGSONG,
    "仿宋": FONT_FANGSONG,
    "楷体_GB2312": FONT_KAI,
    "KaiTi": FONT_KAI,
    "楷体": FONT_KAI,
    "黑体": FONT_HEI,
    "SimHei": FONT_HEI,
    "小标宋": FONT_TITLE,
    "华文中宋": FONT_TITLE,
    "STZhongsong": FONT_TITLE,
}

_DATE_RE = re.compile(r"^\d{4}年\d{1,2}月\d{1,2}日$")
_L1_RE = re.compile(r"^[一二三四五六七八九十百千]+、")
_L2_RE = re.compile(r"^（[一二三四五六七八九十]+）")
_L3_RE = re.compile(r"^\d+\.")
_DRAFT_MARKS = ("讨论稿", "征求意见稿", "送审稿", "草案")


_GONGWEN_FORMAT = ContextVar("gongwen_format_enabled", default=False)


def task_wants_gongwen_format(task: Any) -> bool:
    aid = str(getattr(task, "assistant_id", None) or "")
    skills = [str(x) for x in (getattr(task, "enabled_skill_ids", None) or [])]
    return aid == "official-document-writing" or "official-document-writing" in skills


def enable_gongwen_format(enabled: bool):
    """Return a ContextVar token; reset with `_GONGWEN_FORMAT.reset(token)`."""
    return _GONGWEN_FORMAT.set(bool(enabled))


def reset_gongwen_format(token) -> None:
    if token is not None:
        _GONGWEN_FORMAT.reset(token)


def maybe_apply_gongwen_format(path: str | Path) -> str | None:
    if not _GONGWEN_FORMAT.get():
        return None
    target = Path(path)
    if target.suffix.lower() != ".docx" or not target.is_file():
        return None
    try:
        return apply_body_manuscript_format(target)
    except Exception:
        return None


def apply_body_manuscript_format(path: str | Path) -> str:
    """Rewrite page setup, fonts, line spacing, and page numbers. Return path."""
    target = Path(path)
    doc = Document(str(target))
    _apply_sections(doc)
    _apply_styles(doc)
    _apply_paragraphs(doc)
    _apply_page_numbers(doc)
    doc.save(str(target))
    return str(target)


def _apply_sections(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(3)
        section.bottom_margin = Cm(3)
        section.left_margin = Cm(2.9)
        section.right_margin = Cm(2.9)
        section.header_distance = Cm(1.5)
        section.footer_distance = Cm(1.75)
        section.different_first_page_header_footer = False


def _apply_styles(doc: Document) -> None:
    mapping = {
        "Normal": (FONT_FANGSONG, SIZE_BODY_PT),
        "Title": (FONT_TITLE, SIZE_TITLE_PT),
        "Heading 1": (FONT_HEI, SIZE_BODY_PT),
        "Heading 2": (FONT_KAI, SIZE_BODY_PT),
        "Heading 3": (FONT_FANGSONG, SIZE_BODY_PT),
    }
    for name, (east, size) in mapping.items():
        try:
            style = doc.styles[name]
        except KeyError:
            continue
        _set_element_fonts(style.element, east, size_pt=size)
        pf = style.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(LINE_PT)
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)


def _apply_paragraphs(doc: Document) -> None:
    seen_title = False
    for para in _iter_body_paragraphs(doc):
        text = (para.text or "").strip()
        _set_line_spacing(para)
        kind = _classify(text, seen_title)
        if kind == "title":
            seen_title = True
        east, size, bold, indent = _kind_format(kind)
        if indent:
            _set_first_line_chars(para, 200)
        else:
            _clear_first_line_indent(para)
        if not para.runs and text:
            para.add_run(text)
        for run in para.runs:
            _set_run_fonts(run, east, size_pt=size, bold=bold)


def _classify(text: str, seen_title: bool) -> str:
    if not text:
        return "empty"
    if _DATE_RE.match(text) or text in _DRAFT_MARKS:
        return "date"
    if not seen_title:
        return "title"
    if _L1_RE.match(text):
        return "h1"
    if _L2_RE.match(text):
        return "h2"
    if _L3_RE.match(text):
        return "h3"
    return "body"


def _kind_format(kind: str) -> tuple[str, float, bool | None, bool]:
    if kind == "title":
        return FONT_TITLE, SIZE_TITLE_PT, False, False
    if kind == "date":
        return FONT_KAI, SIZE_BODY_PT, False, False
    if kind == "h1":
        return FONT_HEI, SIZE_BODY_PT, False, True
    if kind == "h2":
        return FONT_KAI, SIZE_BODY_PT, False, True
    if kind == "h3":
        return FONT_FANGSONG, SIZE_BODY_PT, True, True
    if kind == "empty":
        return FONT_FANGSONG, SIZE_BODY_PT, False, False
    return FONT_FANGSONG, SIZE_BODY_PT, False, True


def _set_line_spacing(para) -> None:
    pf = para.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(LINE_PT)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def _set_first_line_chars(para, chars: int) -> None:
    pPr = para._p.get_or_add_pPr()
    ind = pPr.get_or_add_ind()
    ind.set(qn("w:firstLineChars"), str(chars))


def _clear_first_line_indent(para) -> None:
    pPr = para._p.pPr
    if pPr is None or pPr.ind is None:
        return
    ind = pPr.ind
    if qn("w:firstLineChars") in ind.attrib:
        del ind.attrib[qn("w:firstLineChars")]
    ind.firstLine = None


def _set_run_fonts(run, east_asia: str, *, size_pt: float, bold: bool | None) -> None:
    run.font.size = Pt(size_pt)
    run.font.name = FONT_LATIN
    if bold is not None:
        run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    _set_rpr_fonts(rPr, east_asia)


def _set_element_fonts(element, east_asia: str, *, size_pt: float | None = None) -> None:
    rPr = element.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        element.insert(0, rPr)
    if size_pt is not None:
        sz = rPr.find(qn("w:sz"))
        if sz is None:
            sz = OxmlElement("w:sz")
            rPr.append(sz)
        sz.set(qn("w:val"), str(int(size_pt * 2)))
        szCs = rPr.find(qn("w:szCs"))
        if szCs is None:
            szCs = OxmlElement("w:szCs")
            rPr.append(szCs)
        szCs.set(qn("w:val"), str(int(size_pt * 2)))
    _set_rpr_fonts(rPr, east_asia)


def _set_rpr_fonts(rPr, east_asia: str) -> None:
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        try:
            rFonts = rPr.get_or_add_rFonts()
        except Exception:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
    mapped = _EAST_ASIA_REPLACE.get(east_asia, east_asia)
    rFonts.set(qn("w:ascii"), FONT_LATIN)
    rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    rFonts.set(qn("w:eastAsia"), mapped)


def _iter_body_paragraphs(doc: Document):
    for para in doc.paragraphs:
        yield para
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    yield para


def _apply_page_numbers(doc: Document) -> None:
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.clear()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_line_spacing(para)
        left = para.add_run("—— ")
        _set_run_fonts(left, FONT_PAGE, size_pt=SIZE_PAGE_PT, bold=False)
        _append_page_field(para)
        right = para.add_run(" ——")
        _set_run_fonts(right, FONT_PAGE, size_pt=SIZE_PAGE_PT, bold=False)


def _append_page_field(paragraph) -> None:
    def _append_fld_char(kind: str) -> None:
        run = paragraph.add_run()
        _set_run_fonts(run, FONT_PAGE, size_pt=SIZE_PAGE_PT, bold=False)
        node = OxmlElement("w:fldChar")
        node.set(qn("w:fldCharType"), kind)
        run._r.append(node)

    _append_fld_char("begin")
    instr = paragraph.add_run()
    _set_run_fonts(instr, FONT_PAGE, size_pt=SIZE_PAGE_PT, bold=False)
    text = OxmlElement("w:instrText")
    text.set(qn("xml:space"), "preserve")
    text.text = " PAGE "
    instr._r.append(text)
    _append_fld_char("separate")
    cached = paragraph.add_run("1")
    _set_run_fonts(cached, FONT_PAGE, size_pt=SIZE_PAGE_PT, bold=False)
    _append_fld_char("end")
