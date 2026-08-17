from pathlib import Path
import uuid

from langchain_core.tools import BaseTool, tool

from app.guardrails.approval import ConfirmHub
from app.sandbox.path_guard import PathGuard, PathGuardError, _DEFAULT_GUARD, resolve_tool_path, resolve_write_path
from app.workspace.overlay import maybe_record_write

# By default fs tools share the global legacy guard so that callers using
# ``path_guard.set_whitelist()`` continue to work. ``main.py`` injects a
# dedicated ``PathGuard`` instance via ``set_guard()`` in production.
_guard: PathGuard = _DEFAULT_GUARD


def set_guard(guard: PathGuard) -> None:
    """Inject a ``PathGuard`` instance to be used by the fs tools."""
    global _guard
    _guard = guard


def get_guard() -> PathGuard:
    """Return the currently active ``PathGuard`` instance."""
    return _guard


def _prepare_write(path: str) -> Path:
    """Normalize write path (Desktop→workdir remap) then enforce whitelist."""
    from app.sandbox.path_guard import resolve_write_path

    target = resolve_write_path(path)
    _guard.check_path(str(target))
    return target


def _prepare(path: str) -> Path:
    """Normalize aliases then enforce whitelist; return absolute path."""
    target = resolve_tool_path(path)
    _guard.check_path(str(target))
    return target


def _read_docx_text(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = []
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            line = " | ".join(c for c in cells if c)
            if line:
                parts.append(line)
    body = "\n".join(parts).strip()
    if not body:
        return f"[NOTE] {path.name} is a .docx with no extractable paragraphs."
    return body


def _read_file_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        try:
            return _read_docx_text(path)
        except Exception as exc:
            return f"[ERROR] Failed to extract text from docx {path}: {exc}"
    if suffix in {".pptx", ".xlsx", ".xls", ".pdf", ".png", ".jpg", ".jpeg", ".webp", ".gif", ".zip"}:
        return (
            f"[ERROR] {path.name} is a binary/{suffix} file; "
            "fs_read only supports plain text and .docx extraction. "
            "Use a dedicated tool or convert to text first."
        )
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        # Fallback for GBK/legacy office exports without crashing the task.
        raw = path.read_bytes()
        for enc in ("utf-8-sig", "gb18030", "latin-1"):
            try:
                return raw.decode(enc)
            except UnicodeDecodeError:
                continue
        return (
            f"[ERROR] File is not decodable text: {path}. "
            "For Office files prefer .docx (auto-extracted) or convert to UTF-8 text."
        )


@tool
def fs_read(path: str) -> str:
    """Read the text content of a file.

    Prefer absolute paths. ``Desktop`` / ``桌面`` / ``~/Desktop`` map to the
    user's real desktop directory. Relative paths resolve against the frozen
    Space working directory when a task is running.

    ``.docx`` files are text-extracted via python-docx (not raw ZIP bytes).
    """
    try:
        target = _prepare(path)
        return _read_file_text(target)
    except PathGuardError as exc:
        return f"[ERROR] {exc}"
    except OSError as exc:
        return f"[ERROR] {exc}"


@tool
def fs_write(path: str, content: str) -> str:
    """Write text content to a file, creating parent directories as needed.

    Prefer absolute paths under the task working directory. ``Desktop/...``
    aliases are remapped into the workspace when a task is active.
    """
    try:
        target = _prepare_write(path)
        target.parent.mkdir(parents=True, exist_ok=True)
        target.write_text(content, encoding="utf-8")
        maybe_record_write(target)
        return f"Wrote {len(content)} characters to {target}"
    except PathGuardError as exc:
        return f"[ERROR] {exc}"
    except OSError as exc:
        return f"[ERROR] {exc}"


@tool
def fs_list(path: str) -> str:
    """List file names in a directory.

    Prefer absolute paths. ``Desktop`` / ``桌面`` list the user's desktop.
    If *path* points at a file (e.g. ``…/guide.pptx``), list the parent
    directory instead and note that in the response — do not raise.
    """
    try:
        directory = _prepare(path)
    except PathGuardError as exc:
        return f"[ERROR] {exc}"
    note = ""
    if directory.is_file():
        note = (
            f"[NOTE] {directory} is a file, not a directory; "
            f"listing parent {directory.parent} instead.\n"
        )
        directory = directory.parent
    if not directory.is_dir():
        return f"[ERROR] Not a directory: {directory}"
    names = sorted(f.name for f in directory.iterdir() if f.is_file())
    body = "\n".join(names) if names else "(empty)"
    return note + body


def make_fs_write(guard: PathGuard, confirm_hub: ConfirmHub | None) -> BaseTool:
    """Build an ``fs_write`` tool with optional L1 confirmation gate."""

    @tool
    async def fs_write(path: str, content: str) -> str:
        """Write text content to a file, creating parent directories as needed.

        Prefer absolute paths under the task working directory. Desktop aliases
        are remapped into the workspace when a task is active.
        """
        try:
            target = resolve_write_path(path)
            guard.check_path(str(target))
        except PathGuardError as exc:
            return f"[ERROR] {exc}"

        if confirm_hub is not None:
            call_id = f"fs.write:{uuid.uuid4().hex}"
            ok = await confirm_hub.request(
                call_id,
                "fs.write",
                {"path": str(target), "content": content},
            )
            if not ok:
                return "Operation rejected by user"

        target.parent.mkdir(parents=True, exist_ok=True)
        target.write_text(content, encoding="utf-8")
        maybe_record_write(target)
        return f"Wrote {len(content)} characters to {target}"

    return fs_write
