"""Resolve wiki media into text (or a local file path for binaries)."""

from __future__ import annotations

import re
import zipfile
from io import BytesIO
from pathlib import Path
from typing import Any
from urllib.parse import urlparse

from app.tools.builtin.ima.client import ImaClient, ImaError
from app.tools.builtin.ima.notes import get_doc_content
from app.tools.builtin.ima.wiki import get_media_info
from app.tools.builtin.web_fetch import html_to_text
from app.workspace.paths import data_root

MAX_CHARS = 12_000
_SAFE_ID = re.compile(r"[^A-Za-z0-9._-]+")

_EXT_BY_CTYPE = {
    "application/pdf": ".pdf",
    "application/msword": ".doc",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
    "application/vnd.ms-powerpoint": ".ppt",
    "application/vnd.openxmlformats-officedocument.presentationml.presentation": ".pptx",
    "application/vnd.ms-excel": ".xls",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": ".xlsx",
    "text/csv": ".csv",
    "text/markdown": ".md",
    "text/plain": ".txt",
    "text/html": ".html",
    "application/epub+zip": ".epub",
    "image/png": ".png",
    "image/jpeg": ".jpg",
    "image/webp": ".webp",
    "audio/mpeg": ".mp3",
}


def normalize_media_url(url: str) -> str:
    """Turn IMA desktop UNC / protocol-relative paths into https URLs.

    ``get_media_info`` may return ``\\\\res-pkb.ima.qq.com\\...\\file.docx``.
    Opening that as a local file on Windows raises WinError 64.
    """
    raw = (url or "").strip().strip("\"'")
    if not raw:
        return ""
    if raw.lower().startswith(("http://", "https://")):
        return raw
    if raw.lower().startswith("file://"):
        parsed = urlparse(raw)
        host = (parsed.netloc or "").strip()
        path = parsed.path or ""
        if host:
            return f"https://{host}{path}"
        stripped = path.lstrip("/")
        if "ima.qq.com" in stripped.lower():
            return f"https://{stripped}"
        return raw
    if raw.startswith("\\\\") or raw.startswith("//"):
        rest = raw.lstrip("\\/").replace("\\", "/")
        return f"https://{rest}" if rest else ""
    if "ima.qq.com" in raw.lower() and "://" not in raw:
        return "https://" + raw.lstrip("/\\").replace("\\", "/")
    return raw


def _media_type(info: dict[str, Any]) -> int | None:
    raw = info.get("media_type")
    if raw is None:
        raw = info.get("mediaType")
    try:
        return int(raw)
    except (TypeError, ValueError):
        return None


def _looks_like_docx(body: bytes) -> bool:
    if len(body) < 4 or body[:2] != b"PK":
        return False
    try:
        with zipfile.ZipFile(BytesIO(body)) as zf:
            return "word/document.xml" in zf.namelist()
    except Exception:
        return False


def _extract_docx_via_xml(body: bytes) -> str | None:
    """Read w:t runs from document.xml so text boxes still match search."""
    try:
        with zipfile.ZipFile(BytesIO(body)) as zf:
            xml = zf.read("word/document.xml")
    except Exception:
        return None
    texts = re.findall(rb"<w:t[^>]*>([^<]*)</w:t>", xml)
    blob = "".join(t.decode("utf-8", errors="replace") for t in texts).strip()
    return blob or None


def _extract_docx_via_python_docx(body: bytes) -> str | None:
    try:
        from docx import Document
    except ImportError:
        return None
    try:
        doc = Document(BytesIO(body))
    except Exception:
        return None
    parts: list[str] = []
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            line = " | ".join(c for c in cells if c)
            if line:
                parts.append(line)
    try:
        for section in doc.sections:
            for container in (section.header, section.footer):
                for para in container.paragraphs:
                    text = (para.text or "").strip()
                    if text:
                        parts.append(text)
    except Exception:
        pass
    blob = "\n".join(parts).strip()
    return blob or None


def _extract_docx_text(body: bytes) -> str | None:
    via_lib = _extract_docx_via_python_docx(body)
    via_xml = _extract_docx_via_xml(body)
    candidates = [t for t in (via_lib, via_xml) if t]
    if not candidates:
        return None
    return max(candidates, key=len)


def _is_textual(ctype: str, body: bytes) -> bool:
    c = (ctype or "").lower()
    if any(
        token in c
        for token in (
            "text/",
            "json",
            "xml",
            "markdown",
            "html",
            "javascript",
        )
    ):
        return True
    head = body[:24].lstrip().lower()
    return head.startswith((b"<!doctype", b"<html", b"{", b"["))


def _guess_ext(url: str, ctype: str) -> str:
    mapped = _EXT_BY_CTYPE.get((ctype or "").split(";")[0].strip().lower())
    if mapped:
        return mapped
    path = urlparse(url).path
    suffix = Path(path).suffix.lower()
    if suffix and len(suffix) <= 8:
        return suffix
    return ".bin"


def _cache_dir() -> Path:
    path = data_root() / "ima-cache"
    path.mkdir(parents=True, exist_ok=True)
    return path


def _truncate(text: str, max_chars: int) -> str:
    if len(text) <= max_chars:
        return text
    return text[:max_chars] + "\n…(truncated)"


async def load_media_content(
    client: ImaClient,
    media_id: str,
    *,
    max_chars: int = MAX_CHARS,
) -> dict[str, Any]:
    mid = (media_id or "").strip()
    if not mid:
        raise ImaError("media_id 不能为空")
    max_chars = max(500, min(int(max_chars or MAX_CHARS), 40_000))
    info = await get_media_info(client, mid)
    media_type = _media_type(info)

    if media_type == 11:
        notebook = info.get("notebook_ext_info") or {}
        note_id = str(notebook.get("notebook_id") or "").strip()
        if not note_id:
            return {
                "media_id": mid,
                "media_type": media_type,
                "kind": "unavailable",
                "message": "笔记原文不可访问，请使用 IMA 客户端查看。",
            }
        note = await get_doc_content(client, note_id)
        text = _truncate(str(note.get("content") or ""), max_chars)
        return {
            "media_id": mid,
            "media_type": media_type,
            "kind": "note",
            "note_id": note_id,
            "content": text,
            "instruction": "正文已抽出。请据此总结回答用户，不要下载或保存文件。",
        }

    url_info = info.get("url_info") or info.get("urlInfo") or {}
    if not isinstance(url_info, dict):
        url_info = {}
    url = normalize_media_url(str(url_info.get("url") or url_info.get("Url") or ""))
    if not url.lower().startswith(("http://", "https://")):
        return {
            "media_id": mid,
            "media_type": media_type,
            "kind": "unavailable",
            "message": "请使用 IMA 客户端查看原文。",
        }

    raw_headers = url_info.get("headers") or url_info.get("Headers") or {}
    headers = (
        {str(k): str(v) for k, v in raw_headers.items()}
        if isinstance(raw_headers, dict)
        else {}
    )
    try:
        res = await client.fetch_url(url, headers)
        res.raise_for_status()
    except OSError as exc:
        raise ImaError(
            f"拉取原文失败: {exc}。IMA 返回了无法作为本地文件打开的地址，已按 HTTPS 重试仍失败。"
        ) from exc
    except Exception as exc:
        if isinstance(exc, ImaError):
            raise
        raise ImaError(f"拉取原文失败: {exc}") from exc

    ctype = res.headers.get("content-type") or ""
    body = res.content
    ext = _guess_ext(url, ctype)
    # COS often serves docx as application/zip or octet-stream with no .docx in the URL.
    if (
        ext == ".docx"
        or "wordprocessingml" in ctype.lower()
        or media_type == 3
        or _looks_like_docx(body)
    ):
        extracted = _extract_docx_text(body)
        if extracted:
            return {
                "media_id": mid,
                "media_type": media_type,
                "kind": "text",
                "content": _truncate(extracted, max_chars),
                "instruction": "正文已抽出。请据此总结回答用户，不要下载或保存文件。",
            }
        if _looks_like_docx(body):
            ext = ".docx"
    if _is_textual(ctype, body):
        text = body.decode("utf-8", errors="replace")
        if "html" in ctype.lower() or text.lstrip()[:16].lower().startswith(("<!doctype", "<html")):
            text = html_to_text(text)
        return {
            "media_id": mid,
            "media_type": media_type,
            "kind": "text",
            "content": _truncate(text, max_chars),
            "instruction": "正文已抽出。请据此总结回答用户，不要下载或保存文件。",
        }

    safe = _SAFE_ID.sub("_", mid)[:80] or "media"
    dest = _cache_dir() / f"{safe}{ext}"
    dest.write_bytes(body)
    return {
        "media_id": mid,
        "media_type": media_type,
        "kind": "file",
        "url": url,
        "path": str(dest),
        "content_type": ctype,
        "bytes": len(body),
        "message": (
            "无法抽出纯文本。若用户只是搜索，请停在检索摘要，不要下载。"
            "仅当用户要求导出原文件时，才使用本地 path。"
        ),
    }
