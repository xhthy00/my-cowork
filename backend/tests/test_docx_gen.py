"""Unit tests for docx_gen and outline coercion."""

from pathlib import Path

import pytest
from docx import Document
from pydantic import ValidationError

from app.tools.builtin.docgen.coerce import coerce_docx_outline, docx_outline_has_body
from app.tools.builtin.docgen.docx_gen import gen
from app.tools.builtin.docgen.tools import DocxArgs


@pytest.mark.asyncio
async def test_docx_gen_writes_title_and_paragraphs(tmp_path: Path):
    out = tmp_path / "out.docx"
    outline = {
        "title": "Q1 总结",
        "paragraphs": [
            {"heading": "概述", "body": "本季度进展顺利。"},
            {"heading": "指标", "body": "营收增长 12%。"},
            {"heading": "展望", "body": "下一季度聚焦交付。"},
        ],
    }
    path = await gen(outline, str(out))
    assert Path(path).is_file()

    doc = Document(path)
    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    assert texts[0] == "Q1 总结"
    assert "概述" in texts
    assert "本季度进展顺利。" in texts
    assert sum(1 for t in texts if t in ("概述", "指标", "展望")) == 3


@pytest.mark.asyncio
async def test_docx_gen_rejects_title_only(tmp_path: Path):
    out = tmp_path / "empty.docx"
    with pytest.raises(ValueError, match="title only"):
        await gen({"title": "成绩最佳学生个人成绩分析报告", "paragraphs": []}, str(out))
    assert not out.exists()


def test_coerce_accepts_sections_content_alias():
    outline = coerce_docx_outline(
        {
            "title": "报告",
            "sections": [
                {"title": "核心发现", "content": "总分 144，年级第1。"},
                {"heading": "建议", "text": "巩固填空题。"},
            ],
        }
    )
    assert docx_outline_has_body(outline)
    assert outline.paragraphs[0].heading == "核心发现"
    assert "144" in outline.paragraphs[0].body


def test_coerce_nested_sections_paragraphs_and_tables():
    """LLM often sends sections[].paragraphs[] instead of paragraphs[].body."""
    outline = coerce_docx_outline(
        {
            "title": "最后一名与中位数差距分析报告",
            "sections": [
                {
                    "heading": "一、核心数据",
                    "paragraphs": [
                        "最后一名总分 78 分。",
                        "班级中位数为 126 分。",
                    ],
                    "callout": "差距达 48 分。",
                    "tables": [
                        {
                            "caption": "表1",
                            "headers": ["指标", "分值"],
                            "rows": [["最后一名", "78"], ["中位数", "126"]],
                        }
                    ],
                    "subsections": [
                        {
                            "heading": "1. 题型拆解",
                            "paragraphs": ["选择题失分最多。"],
                        }
                    ],
                }
            ],
        }
    )
    assert docx_outline_has_body(outline)
    bodies = "\n".join(p.body for p in outline.paragraphs)
    headings = [p.heading for p in outline.paragraphs]
    assert "一、核心数据" in headings
    assert "1. 题型拆解" in headings
    assert "78" in bodies and "126" in bodies
    assert "选择题失分最多" in bodies
    assert "表1" in bodies


def test_coerce_accepts_top_level_content_and_item_wrapper():
    outline = coerce_docx_outline(
        {
            "title": "报告",
            "paragraphs": {"item": [{"heading": "A", "body": "正文A"}]},
        }
    )
    assert outline.paragraphs[0].body == "正文A"

    outline2 = coerce_docx_outline({"title": "报告", "content": "整篇正文写这里。"})
    assert outline2.paragraphs[0].body == "整篇正文写这里。"


def test_docx_args_rejects_empty_body():
    with pytest.raises(ValidationError):
        DocxArgs.model_validate(
            {
                "outline": {"title": "只有标题", "paragraphs": []},
                "out_path": "/tmp/x.docx",
            }
        )
