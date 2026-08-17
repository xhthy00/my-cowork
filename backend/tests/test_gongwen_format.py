"""公文正文稿 formatter overwrites GB/T 9704 page setup."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from app.tools.builtin.docgen.gongwen_format import (
    FONT_FANGSONG,
    FONT_TITLE,
    apply_body_manuscript_format,
    enable_gongwen_format,
    maybe_apply_gongwen_format,
    reset_gongwen_format,
    task_wants_gongwen_format,
)


def _east_asia(run) -> str | None:
    rPr = run._element.rPr
    if rPr is None or rPr.rFonts is None:
        return None
    return rPr.rFonts.get(qn("w:eastAsia"))


def test_apply_overwrites_gb_t_9704_page_setup(tmp_path: Path):
    path = tmp_path / "gongwen.docx"
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)
    title = doc.add_paragraph("关于增加项目经费的请示")
    title.runs[0].font.name = "仿宋_GB2312"
    title.runs[0].font.size = Pt(16)
    body = doc.add_paragraph("根据有关规定，现将有关事项请示如下。")
    run = body.runs[0]
    run.font.name = "仿宋_GB2312"
    run.font.size = Pt(16)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")
    body.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    body.paragraph_format.line_spacing = Pt(28)
    doc.save(path)

    apply_body_manuscript_format(path)

    out = Document(str(path))
    sec = out.sections[0]
    assert abs(sec.top_margin.cm - 3.0) < 0.05
    assert abs(sec.bottom_margin.cm - 3.0) < 0.05
    assert abs(sec.left_margin.cm - 2.9) < 0.05
    assert abs(sec.right_margin.cm - 2.9) < 0.05
    assert abs(sec.header_distance.cm - 1.5) < 0.05
    assert abs(sec.footer_distance.cm - 1.75) < 0.05

    body_para = [p for p in out.paragraphs if "请示如下" in p.text][0]
    assert abs(float(body_para.paragraph_format.line_spacing.pt) - 29) < 0.2
    assert _east_asia(body_para.runs[0]) == FONT_FANGSONG

    title_para = [p for p in out.paragraphs if p.text.startswith("关于")][0]
    assert _east_asia(title_para.runs[0]) == FONT_TITLE

    footer_xml = out.sections[0].footer._element.xml
    assert "PAGE" in footer_xml
    assert "——" in out.sections[0].footer.paragraphs[0].text or "PAGE" in footer_xml


def test_maybe_apply_respects_context(tmp_path: Path):
    path = tmp_path / "a.docx"
    Document().save(path)
    assert maybe_apply_gongwen_format(path) is None
    token = enable_gongwen_format(True)
    try:
        assert maybe_apply_gongwen_format(path) == str(path)
    finally:
        reset_gongwen_format(token)


def test_task_wants_gongwen_format():
    class T:
        assistant_id = "official-document-writing"
        enabled_skill_ids = []

    assert task_wants_gongwen_format(T()) is True

    class T2:
        assistant_id = "word-creator"
        enabled_skill_ids = ["officecli-docx"]

    assert task_wants_gongwen_format(T2()) is False
