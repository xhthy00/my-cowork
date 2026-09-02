"""IMA OpenAPI client, tools, and content loading."""

from __future__ import annotations

import json

import httpx
import pytest
from fastapi.testclient import TestClient

from app.main import create_app
from app.tools.builtin.ima.client import ImaClient, ImaError
from app.tools.builtin.ima.content import load_media_content, normalize_media_url
from app.tools.builtin.ima.credentials import (
    MISSING_CREDENTIALS_MSG,
    load_credentials,
)
from app.tools.builtin.ima.tools import (
    ima_get_media_content,
    ima_list_knowledge_bases,
    ima_search_knowledge,
    make_ima_tools,
)
from app.tools.builtin.ima.wiki import get_media_info, search_knowledge_base


def _clear_creds(monkeypatch: pytest.MonkeyPatch) -> None:
    for key in (
        "IMA_OPENAPI_CLIENTID",
        "IMA_OPENAPI_APIKEY",
        "IMA_CLIENT_ID",
        "IMA_API_KEY",
        "IMA_BASE_URL",
        "MY_COWORK_CONFIG",
    ):
        monkeypatch.delenv(key, raising=False)
    monkeypatch.setattr("app.tools.builtin.ima.credentials._toml_ima", lambda: {})
    monkeypatch.setattr("app.tools.builtin.ima.credentials._read_file", lambda _p: "")


def _set_creds(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setenv("IMA_OPENAPI_CLIENTID", "cid-test")
    monkeypatch.setenv("IMA_OPENAPI_APIKEY", "key-test")


def _client_for(handler) -> ImaClient:
    from app.tools.builtin.ima.credentials import ImaCredentials

    return ImaClient(
        ImaCredentials("cid-test", "key-test"),
        transport=httpx.MockTransport(handler),
    )


def test_missing_credentials_message(monkeypatch: pytest.MonkeyPatch) -> None:
    _clear_creds(monkeypatch)
    assert load_credentials() is None


def test_env_credentials_win(monkeypatch: pytest.MonkeyPatch) -> None:
    _clear_creds(monkeypatch)
    _set_creds(monkeypatch)
    creds = load_credentials()
    assert creds is not None
    assert creds.client_id == "cid-test"
    assert creds.api_key == "key-test"


@pytest.mark.asyncio
async def test_list_bases_parses_info_list(monkeypatch: pytest.MonkeyPatch) -> None:
    _set_creds(monkeypatch)

    def handler(request: httpx.Request) -> httpx.Response:
        assert request.headers.get("ima-openapi-clientid") == "cid-test"
        assert request.headers.get("ima-openapi-ctx") == "skill_version=1.1.9"
        assert request.url.path.endswith("/openapi/wiki/v1/search_knowledge_base")
        body = json.loads(request.content.decode())
        assert body["query"] == "产品"
        return httpx.Response(
            200,
            json={
                "code": 0,
                "msg": "ok",
                "data": {
                    "info_list": [{"id": "kb1", "name": "产品文档", "cover_url": ""}],
                    "is_end": True,
                    "next_cursor": "",
                },
            },
        )

    monkeypatch.setattr(
        "app.tools.builtin.ima.tools.ImaClient",
        lambda: _client_for(handler),
    )
    blob = await ima_list_knowledge_bases(query="产品")
    data = json.loads(blob)
    assert data["items"][0]["name"] == "产品文档"
    assert data["items"][0]["id"] == "kb1"


@pytest.mark.asyncio
async def test_list_bases_falls_back_to_addable(monkeypatch: pytest.MonkeyPatch) -> None:
    _set_creds(monkeypatch)

    def handler(request: httpx.Request) -> httpx.Response:
        if request.url.path.endswith("/openapi/wiki/v1/search_knowledge_base"):
            return httpx.Response(
                200,
                json={"code": 0, "data": {"info_list": [], "is_end": True}},
            )
        assert request.url.path.endswith(
            "/openapi/wiki/v1/get_addable_knowledge_base_list"
        )
        return httpx.Response(
            200,
            json={
                "code": 0,
                "data": {
                    "addable_knowledge_base_list": [
                        {"id": "kb-add", "name": "可写入库"},
                    ],
                    "is_end": True,
                    "next_cursor": "",
                },
            },
        )

    monkeypatch.setattr(
        "app.tools.builtin.ima.tools.ImaClient",
        lambda: _client_for(handler),
    )
    blob = await ima_list_knowledge_bases()
    data = json.loads(blob)
    assert data["items"][0]["id"] == "kb-add"
    assert "hint" not in data


@pytest.mark.asyncio
async def test_list_bases_empty_is_not_missing_credentials(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    _set_creds(monkeypatch)

    def handler(request: httpx.Request) -> httpx.Response:
        return httpx.Response(
            200,
            json={"code": 0, "data": {"info_list": [], "is_end": True}},
        )

    monkeypatch.setattr(
        "app.tools.builtin.ima.tools.ImaClient",
        lambda: _client_for(handler),
    )
    blob = await ima_list_knowledge_bases()
    data = json.loads(blob)
    assert data["items"] == []
    assert "凭证已生效" in data["hint"]
    assert "同一账号" in data["hint"]
    assert not blob.startswith("[ERROR]")


@pytest.mark.asyncio
async def test_list_bases_accepts_kb_id_alias(monkeypatch: pytest.MonkeyPatch) -> None:
    _set_creds(monkeypatch)

    def handler(request: httpx.Request) -> httpx.Response:
        return httpx.Response(
            200,
            json={
                "code": 0,
                "data": {
                    "info_list": [{"kb_id": "kb2", "kb_name": "别名库"}],
                    "is_end": True,
                },
            },
        )

    monkeypatch.setattr(
        "app.tools.builtin.ima.tools.ImaClient",
        lambda: _client_for(handler),
    )
    blob = await ima_list_knowledge_bases(query="别名")
    data = json.loads(blob)
    assert data["items"][0] == {"id": "kb2", "name": "别名库", "cover_url": ""}


def test_extract_camel_case_info_list() -> None:
    from app.tools.builtin.ima.wiki import extract_knowledge_bases

    rows, is_end, _, keys = extract_knowledge_bases(
        {"infoList": [{"kbId": "kb9", "kbName": "驼峰库"}], "isEnd": True}
    )
    assert rows[0]["id"] == "kb9"
    assert rows[0]["name"] == "驼峰库"
    assert is_end is True
    assert "infoList" in keys


@pytest.mark.asyncio
async def test_search_knowledge_returns_highlights(monkeypatch: pytest.MonkeyPatch) -> None:
    _set_creds(monkeypatch)

    def handler(request: httpx.Request) -> httpx.Response:
        assert request.url.path.endswith("/openapi/wiki/v1/search_knowledge")
        return httpx.Response(
            200,
            json={
                "code": 0,
                "data": {
                    "info_list": [
                        {
                            "media_id": "m1",
                            "title": "退款政策",
                            "parent_folder_id": "",
                            "highlight_content": "7天无理由",
                        }
                    ],
                    "is_end": True,
                    "next_cursor": "",
                },
            },
        )

    monkeypatch.setattr(
        "app.tools.builtin.ima.tools.ImaClient",
        lambda: _client_for(handler),
    )
    blob = await ima_search_knowledge("退款", "kb1")
    data = json.loads(blob)
    assert data["items"][0]["highlight_content"] == "7天无理由"
    assert "summarize" in data["note"]
    assert "Do not download" in data["note"]
    assert "ima_get_media_content" in data["note"]


def test_split_search_terms_strips_slashes() -> None:
    from app.tools.builtin.ima.tools import split_search_terms

    assert split_search_terms("涉密 / 系统集成 / 资质") == ["涉密", "系统集成", "资质"]
    assert split_search_terms("总体集成") == ["总体集成"]


@pytest.mark.asyncio
async def test_search_splits_slash_query(monkeypatch: pytest.MonkeyPatch) -> None:
    _set_creds(monkeypatch)
    seen: list[str] = []

    def handler(request: httpx.Request) -> httpx.Response:
        if request.url.path.endswith("/search_knowledge"):
            body = json.loads(request.content.decode())
            seen.append(body["query"])
            title = "涉密系统资质要求.docx" if body["query"] == "涉密" else ""
            items = (
                [{"media_id": "m-secret", "title": title, "highlight_content": ""}]
                if title
                else []
            )
            return httpx.Response(
                200,
                json={"code": 0, "data": {"info_list": items, "is_end": True}},
            )
        return httpx.Response(200, json={"code": 0, "data": {}})

    monkeypatch.setattr(
        "app.tools.builtin.ima.tools.ImaClient",
        lambda: _client_for(handler),
    )
    blob = await ima_search_knowledge("涉密 / 系统集成 / 资质", "kb1")
    data = json.loads(blob)
    assert seen == ["涉密", "系统集成", "资质"]
    assert data["items"][0]["title"] == "涉密系统资质要求.docx"
    assert data["terms"] == ["涉密", "系统集成", "资质"]


@pytest.mark.asyncio
async def test_business_error_surfaces_msg(monkeypatch: pytest.MonkeyPatch) -> None:
    _set_creds(monkeypatch)

    def handler(_request: httpx.Request) -> httpx.Response:
        return httpx.Response(200, json={"code": 20004, "msg": "鉴权失败"})

    client = _client_for(handler)
    with pytest.raises(ImaError, match="鉴权失败") as exc:
        await search_knowledge_base(client, query="")
    assert exc.value.code == 20004


@pytest.mark.asyncio
async def test_missing_credentials_tool_error(monkeypatch: pytest.MonkeyPatch) -> None:
    _clear_creds(monkeypatch)
    blob = await ima_list_knowledge_bases()
    assert blob.startswith("[ERROR]")
    assert "知识库" in blob
    assert "cid-test" not in blob
    assert "key-test" not in blob
    assert MISSING_CREDENTIALS_MSG.split("。")[0] in blob


@pytest.mark.asyncio
async def test_get_media_info_url_text(monkeypatch: pytest.MonkeyPatch) -> None:
    def handler(request: httpx.Request) -> httpx.Response:
        if request.method == "POST":
            return httpx.Response(
                200,
                json={
                    "code": 0,
                    "data": {
                        "media_type": 7,
                        "url_info": {
                            "url": "https://cdn.example/doc.md",
                            "headers": {"Authorization": "Bearer tok"},
                        },
                    },
                },
            )
        assert request.headers.get("Authorization") == "Bearer tok"
        assert "ima-openapi-apikey" not in {k.lower() for k in request.headers.keys()}
        return httpx.Response(200, text="# hello ima", headers={"content-type": "text/markdown"})

    client = _client_for(handler)
    out = await load_media_content(client, "media-1")
    assert out["kind"] == "text"
    assert "hello ima" in out["content"]


@pytest.mark.asyncio
async def test_get_media_note_branch() -> None:
    def handler(request: httpx.Request) -> httpx.Response:
        path = request.url.path
        if path.endswith("/get_media_info"):
            return httpx.Response(
                200,
                json={
                    "code": 0,
                    "data": {
                        "media_type": 11,
                        "notebook_ext_info": {"notebook_id": "note-9"},
                    },
                },
            )
        assert path.endswith("/get_doc_content")
        body = json.loads(request.content.decode())
        assert body["note_id"] == "note-9"
        assert body["target_content_format"] == 0
        return httpx.Response(200, json={"code": 0, "data": {"content": "笔记正文"}})

    out = await load_media_content(_client_for(handler), "media-note")
    assert out["kind"] == "note"
    assert out["content"] == "笔记正文"


def test_normalize_unc_media_url() -> None:
    unc = (
        r"\\res-pkb.ima.qq.com\2\eWDaG8YFwTyPDl1e8fofxe\file_manager"
        r"\d26617ceab6730e004ad.docx"
    )
    assert normalize_media_url(unc) == (
        "https://res-pkb.ima.qq.com/2/eWDaG8YFwTyPDl1e8fofxe/file_manager/"
        "d26617ceab6730e004ad.docx"
    )
    assert (
        normalize_media_url("//res-pkb.ima.qq.com/2/a.docx")
        == "https://res-pkb.ima.qq.com/2/a.docx"
    )
    assert normalize_media_url("https://cdn.example/a.md") == "https://cdn.example/a.md"


@pytest.mark.asyncio
async def test_unc_docx_is_fetched_over_https_and_extracted() -> None:
    from io import BytesIO

    from docx import Document

    buf = BytesIO()
    doc = Document()
    doc.add_paragraph("知识库正文段落")
    doc.save(buf)
    payload = buf.getvalue()

    def handler(request: httpx.Request) -> httpx.Response:
        if request.method == "POST":
            return httpx.Response(
                200,
                json={
                    "code": 0,
                    "data": {
                        "media_type": 3,
                        "url_info": {
                            "url": r"\\res-pkb.ima.qq.com\2\abc\file_manager\x.docx",
                            "headers": {"Authorization": "Bearer tok"},
                        },
                    },
                },
            )
        assert str(request.url).startswith("https://res-pkb.ima.qq.com/")
        assert request.url.path.endswith(".docx")
        return httpx.Response(
            200,
            content=payload,
            headers={
                "content-type": (
                    "application/vnd.openxmlformats-officedocument."
                    "wordprocessingml.document"
                )
            },
        )

    out = await load_media_content(_client_for(handler), "media-docx")
    assert out["kind"] == "text"
    assert "知识库正文段落" in out["content"]
    assert "url" not in out
    assert "不要下载" in out["instruction"]


def _docx_bytes(paragraph: str) -> bytes:
    from io import BytesIO

    from docx import Document

    buf = BytesIO()
    doc = Document()
    doc.add_paragraph(paragraph)
    doc.save(buf)
    return buf.getvalue()


@pytest.mark.asyncio
async def test_docx_octet_stream_without_extension_is_extracted() -> None:
    payload = _docx_bytes("总体集成资质要求")

    def handler(request: httpx.Request) -> httpx.Response:
        if request.method == "POST":
            return httpx.Response(
                200,
                json={
                    "code": 0,
                    "data": {
                        "mediaType": "3",
                        "urlInfo": {
                            "url": "https://res-pkb.ima.qq.com/2/abc/file_manager/hash",
                            "headers": {},
                        },
                    },
                },
            )
        return httpx.Response(
            200,
            content=payload,
            headers={"content-type": "application/octet-stream"},
        )

    out = await load_media_content(_client_for(handler), "media-zip")
    assert out["kind"] == "text"
    assert "总体集成" in out["content"]


@pytest.mark.asyncio
async def test_search_falls_back_to_docx_body(monkeypatch: pytest.MonkeyPatch) -> None:
    _set_creds(monkeypatch)
    payload = _docx_bytes("项目范围包括总体集成与运维保障。")

    def handler(request: httpx.Request) -> httpx.Response:
        path = request.url.path
        if path.endswith("/search_knowledge"):
            return httpx.Response(
                200,
                json={"code": 0, "data": {"info_list": [], "is_end": True}},
            )
        if path.endswith("/get_knowledge_list"):
            return httpx.Response(
                200,
                json={
                    "code": 0,
                    "data": {
                        "knowledge_list": [
                            {
                                "media_id": "m-docx",
                                "title": "涉密系统资质要求.docx",
                            }
                        ],
                        "is_end": True,
                    },
                },
            )
        if path.endswith("/get_media_info"):
            return httpx.Response(
                200,
                json={
                    "code": 0,
                    "data": {
                        "media_type": 3,
                        "url_info": {
                            "url": "https://cdn.example/hash",
                            "headers": {},
                        },
                    },
                },
            )
        return httpx.Response(
            200,
            content=payload,
            headers={"content-type": "application/zip"},
        )

    monkeypatch.setattr(
        "app.tools.builtin.ima.tools.ImaClient",
        lambda: _client_for(handler),
    )
    blob = await ima_search_knowledge("总体集成", "kb1")
    data = json.loads(blob)
    assert data["source"] == "body"
    assert data["items"][0]["title"] == "涉密系统资质要求.docx"
    assert "总体集成" in data["items"][0]["highlight_content"]


@pytest.mark.asyncio
async def test_search_keeps_folder_hits(monkeypatch: pytest.MonkeyPatch) -> None:
    _set_creds(monkeypatch)

    def handler(request: httpx.Request) -> httpx.Response:
        assert request.url.path.endswith("/search_knowledge")
        return httpx.Response(
            200,
            json={
                "code": 0,
                "data": {
                    "infoList": [
                        {
                            "folderId": "folder_overall",
                            "name": "总体集成",
                            "parentFolderId": "",
                        }
                    ],
                    "isEnd": True,
                },
            },
        )

    monkeypatch.setattr(
        "app.tools.builtin.ima.tools.ImaClient",
        lambda: _client_for(handler),
    )
    blob = await ima_search_knowledge("总体集成", "kb1")
    data = json.loads(blob)
    assert data["items"][0]["kind"] == "folder"
    assert data["items"][0]["title"] == "总体集成"
    assert data["items"][0]["media_id"] == "folder_overall"


@pytest.mark.asyncio
async def test_body_fallback_enters_subfolder(monkeypatch: pytest.MonkeyPatch) -> None:
    _set_creds(monkeypatch)
    payload = _docx_bytes("项目范围包括总体集成与运维保障。")

    def handler(request: httpx.Request) -> httpx.Response:
        path = request.url.path
        if path.endswith("/search_knowledge"):
            return httpx.Response(
                200,
                json={"code": 0, "data": {"info_list": [], "is_end": True}},
            )
        if path.endswith("/get_knowledge_list"):
            body = json.loads(request.content.decode())
            if body.get("folder_id") == "folder_sub":
                items = [
                    {"media_id": "m-nested", "title": "总体方案.docx"},
                ]
            else:
                items = [
                    {"folder_id": "folder_sub", "name": "集成资料"},
                    {"media_id": "m-root", "title": "itss.docx"},
                ]
            return httpx.Response(
                200,
                json={
                    "code": 0,
                    "data": {"knowledge_list": items, "is_end": True},
                },
            )
        if path.endswith("/get_media_info"):
            return httpx.Response(
                200,
                json={
                    "code": 0,
                    "data": {
                        "media_type": 3,
                        "url_info": {"url": "https://cdn.example/hash", "headers": {}},
                    },
                },
            )
        return httpx.Response(
            200,
            content=payload,
            headers={"content-type": "application/zip"},
        )

    monkeypatch.setattr(
        "app.tools.builtin.ima.tools.ImaClient",
        lambda: _client_for(handler),
    )
    blob = await ima_search_knowledge("总体集成", "kb1")
    data = json.loads(blob)
    assert data["source"] == "body"
    titles = {item["title"] for item in data["items"]}
    assert "总体方案.docx" in titles


@pytest.mark.asyncio
async def test_get_media_unavailable() -> None:
    def handler(_request: httpx.Request) -> httpx.Response:
        return httpx.Response(200, json={"code": 0, "data": {"media_type": 1}})

    out = await load_media_content(_client_for(handler), "media-x")
    assert out["kind"] == "unavailable"
    assert "IMA 客户端" in out["message"]


@pytest.mark.asyncio
async def test_binary_media_saves_file(tmp_path, monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr("app.tools.builtin.ima.content.data_root", lambda: tmp_path)

    def handler(request: httpx.Request) -> httpx.Response:
        if request.method == "POST":
            return httpx.Response(
                200,
                json={
                    "code": 0,
                    "data": {
                        "media_type": 1,
                        "url_info": {"url": "https://cdn.example/a.pdf", "headers": {}},
                    },
                },
            )
        return httpx.Response(
            200,
            content=b"%PDF-1.4 fake",
            headers={"content-type": "application/pdf"},
        )

    out = await load_media_content(_client_for(handler), "pdf-1")
    assert out["kind"] == "file"
    assert out["path"].endswith(".pdf")
    assert (tmp_path / "ima-cache").exists()


@pytest.mark.asyncio
async def test_ima_get_media_content_tool(monkeypatch: pytest.MonkeyPatch) -> None:
    _set_creds(monkeypatch)

    def handler(request: httpx.Request) -> httpx.Response:
        if request.url.path.endswith("/get_media_info"):
            return httpx.Response(
                200,
                json={"code": 0, "data": {"media_type": 1}},
            )
        return httpx.Response(200, json={"code": 0, "data": {}})

    monkeypatch.setattr(
        "app.tools.builtin.ima.tools.ImaClient",
        lambda: _client_for(handler),
    )
    blob = await ima_get_media_content("m-empty")
    data = json.loads(blob)
    assert data["kind"] == "unavailable"


def test_make_ima_tools_names() -> None:
    names = [t.name for t in make_ima_tools()]
    assert names == [
        "ima_list_knowledge_bases",
        "ima_get_knowledge_base",
        "ima_list_knowledge",
        "ima_search_knowledge",
        "ima_get_media_content",
    ]


def test_ima_status_and_test_routes(monkeypatch: pytest.MonkeyPatch) -> None:
    from unittest.mock import MagicMock

    _clear_creds(monkeypatch)
    app = create_app(task_manager=MagicMock(), bus=None, confirm_hub=MagicMock())
    client = TestClient(app)
    res = client.get("/api/ima/status")
    assert res.status_code == 200
    assert res.json()["configured"] is False

    listed = client.get("/api/ima/knowledge-bases")
    assert listed.status_code == 200
    assert listed.json() == {
        "configured": False,
        "items": [],
        "empty": True,
        "hint": "",
    }

    missing = client.post("/api/ima/test")
    assert missing.status_code == 400
    assert "知识库" in missing.json()["detail"]

    _set_creds(monkeypatch)

    async def fake_search(_client, **_kwargs):
        return {"items": [{"id": "kb1", "name": "库"}]}

    monkeypatch.setattr(
        "app.server.routes.ima.list_visible_knowledge_bases", fake_search
    )
    ok = client.post("/api/ima/test")
    assert ok.status_code == 200
    assert ok.json()["ok"] is True
    assert ok.json()["sample_count"] == 1

    listed = client.get("/api/ima/knowledge-bases")
    assert listed.status_code == 200
    body = listed.json()
    assert body["configured"] is True
    assert body["items"] == [{"id": "kb1", "name": "库", "source": "ima"}]


def test_ima_knowledge_bases_upstream_error_is_not_unconfigured(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from unittest.mock import MagicMock

    from app.tools.builtin.ima.client import ImaError

    _set_creds(monkeypatch)
    app = create_app(task_manager=MagicMock(), bus=None, confirm_hub=MagicMock())
    client = TestClient(app)

    async def boom(_client, **_kwargs):
        raise ImaError("ima down")

    monkeypatch.setattr(
        "app.server.routes.ima.list_visible_knowledge_bases", boom
    )
    listed = client.get("/api/ima/knowledge-bases")
    assert listed.status_code == 200
    body = listed.json()
    assert body["configured"] is True
    assert body["items"] == []
    assert "ima down" in body["hint"]


@pytest.mark.asyncio
async def test_get_media_info_passthrough() -> None:
    def handler(_request: httpx.Request) -> httpx.Response:
        return httpx.Response(200, json={"code": 0, "data": {"media_type": 2}})

    data = await get_media_info(_client_for(handler), "m1")
    assert data["media_type"] == 2
