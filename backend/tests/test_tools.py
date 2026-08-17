from pathlib import Path

import pytest
from langchain_core.tools import tool as langchain_tool

from app.sandbox.path_guard import PathGuard
from app.tools.builtin.fs import fs_list, fs_read, fs_write, set_guard
from app.tools.registry import ToolRegistry


class TestToolRegistry:
    def test_register_and_list_tools(self):
        registry = ToolRegistry()

        @langchain_tool
        def mock_tool(query: str) -> str:
            """A mock tool."""
            return f"mock:{query}"

        registry.register("mock.tool", mock_tool)
        tools = registry.list_tools()
        assert len(tools) == 1
        assert tools[0].name == "mock_tool"

    def test_get_tool_by_name(self):
        registry = ToolRegistry()

        @langchain_tool
        def mock_tool(query: str) -> str:
            """A mock tool."""
            return f"mock:{query}"

        registry.register("mock.tool", mock_tool)
        assert registry.get("mock.tool") is mock_tool

    def test_get_missing_tool_raises(self):
        registry = ToolRegistry()
        with pytest.raises(KeyError):
            registry.get("missing.tool")


class TestFsTools:
    @pytest.fixture(autouse=True)
    def _whitelist(self, tmp_path):
        set_guard(PathGuard([str(tmp_path)]))

    def test_write_and_read(self, tmp_path):
        path = str(tmp_path / "hello.txt")
        fs_write.invoke({"path": path, "content": "hi"})
        assert fs_read.invoke({"path": path}) == "hi"

    def test_list_directory(self, tmp_path):
        (tmp_path / "a.txt").write_text("a")
        (tmp_path / "b.txt").write_text("b")
        result = fs_list.invoke({"path": str(tmp_path)})
        names = sorted(line for line in str(result).splitlines() if line.strip())
        assert names == ["a.txt", "b.txt"]

    def test_list_file_path_lists_parent(self, tmp_path):
        f = tmp_path / "guide.pptx"
        f.write_bytes(b"PK")
        result = str(fs_list.invoke({"path": str(f)}))
        assert "is a file" in result
        assert "guide.pptx" in result
        # Must not raise / crash the agent turn.

    def test_write_outside_whitelist_returns_error(self, tmp_path):
        outside = "/tmp/outside.txt"
        result = str(fs_write.invoke({"path": outside, "content": "nope"}))
        assert result.startswith("[ERROR]")
        assert "whitelist" in result.lower()

    def test_read_outside_whitelist_returns_error(self, tmp_path):
        outside = "/tmp/outside.txt"
        result = str(fs_read.invoke({"path": outside}))
        assert result.startswith("[ERROR]")
        assert "whitelist" in result.lower()

    def test_list_root_returns_error_not_crash(self, tmp_path):
        result = str(fs_list.invoke({"path": "/"}))
        assert result.startswith("[ERROR]")
        assert "whitelist" in result.lower()

    def test_write_creates_parent_directories(self, tmp_path):
        path = str(tmp_path / "subdir" / "nested.txt")
        fs_write.invoke({"path": path, "content": "nested"})
        assert Path(path).read_text() == "nested"

    def test_read_docx_extracts_text(self, tmp_path):
        from docx import Document

        path = tmp_path / "note.docx"
        doc = Document()
        doc.add_paragraph("你好世界")
        doc.add_paragraph("second line")
        doc.save(str(path))
        text = str(fs_read.invoke({"path": str(path)}))
        assert "你好世界" in text
        assert "second line" in text
        assert not text.startswith("[ERROR]")

    def test_read_gbk_text_does_not_crash(self, tmp_path):
        path = tmp_path / "gbk.txt"
        path.write_bytes("中文内容".encode("gbk"))
        text = str(fs_read.invoke({"path": str(path)}))
        assert "中文内容" in text

    def test_read_pptx_returns_error_not_crash(self, tmp_path):
        path = tmp_path / "deck.pptx"
        path.write_bytes(b"PK\x03\x04" + b"\xe4" * 100)
        result = str(fs_read.invoke({"path": str(path)}))
        assert result.startswith("[ERROR]")
        assert "pptx" in result.lower() or "binary" in result.lower()
